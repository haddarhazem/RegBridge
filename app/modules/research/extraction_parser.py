"""Deterministic text parsing and locators for exact research versions."""

from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(frozen=True)
class SourceLocator:
    document_version_id: str
    locator_type: str
    paragraph: int
    start_char: int
    end_char: int


@dataclass(frozen=True)
class SourceSegment:
    segment_id: str
    document_version_id: str
    locator: SourceLocator
    text: str


@dataclass(frozen=True)
class ParsedSource:
    document_version_id: str
    mime_type: str
    text: str
    paragraphs: tuple[str, ...]

    def locator(self, paragraph: int) -> SourceLocator:
        if paragraph < 0 or paragraph >= len(self.paragraphs):
            raise ValueError("paragraph is outside the parsed source")
        start = sum(len(item) + 1 for item in self.paragraphs[:paragraph])
        return SourceLocator(self.document_version_id, "paragraph", paragraph, start, start + len(self.paragraphs[paragraph]))

    def resolve(self, locator: dict) -> str:
        if str(locator.get("document_version_id")) != self.document_version_id:
            raise ValueError("locator belongs to another document version")
        if locator.get("locator_type") != "paragraph":
            raise ValueError("unsupported locator type")
        paragraph = int(locator.get("paragraph", -1))
        expected = self.locator(paragraph)
        start = int(locator.get("start_char", expected.start_char))
        end = int(locator.get("end_char", expected.end_char))
        if start != expected.start_char or end != expected.end_char or start < 0 or end > len(self.text) or start >= end:
            raise ValueError("locator bounds are invalid")
        value = self.text[start:end]
        if not value.strip():
            raise ValueError("locator resolves empty text")
        return value


def parse_source(document_version_id: str, mime_type: str, payload: bytes) -> ParsedSource:
    if not payload:
        raise ValueError("SOURCE_TEXT_UNAVAILABLE: empty source")
    if mime_type == "text/plain":
        raw_text = payload.decode("utf-8")
        text = "\n".join(raw_text.splitlines())
        paragraphs = tuple(text.splitlines() or [text])
    elif mime_type == "application/pdf":
        try:
            pages = [page.extract_text() or "" for page in PdfReader(io.BytesIO(payload)).pages]
        except Exception as exc:
            raise ValueError("SOURCE_TEXT_UNAVAILABLE: PDF text extraction failed") from exc
        if not any(page.strip() for page in pages):
            raise ValueError("SOURCE_TEXT_UNAVAILABLE: PDF contains no extractable text")
        paragraphs = tuple(line for page in pages for line in page.splitlines()) or ("",)
        text = "\n".join(paragraphs)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            document = DocxDocument(io.BytesIO(payload))
            paragraphs = tuple(paragraph.text for paragraph in document.paragraphs)
        except Exception as exc:
            raise ValueError("SOURCE_TEXT_UNAVAILABLE: DOCX text extraction failed") from exc
        if not any(item.strip() for item in paragraphs):
            raise ValueError("SOURCE_TEXT_UNAVAILABLE: DOCX contains no extractable text")
        text = "\n".join(paragraphs)
    else:
        raise ValueError("SOURCE_TEXT_UNAVAILABLE: unsupported source MIME type")
    if not text.strip():
        raise ValueError("SOURCE_TEXT_UNAVAILABLE: source contains no text")
    return ParsedSource(document_version_id, mime_type, text, paragraphs)


def segment_source(parsed: ParsedSource) -> tuple[SourceSegment, ...]:
    """Create an allowlisted, deterministic segment namespace for one source."""
    segments: list[SourceSegment] = []
    for index, value in enumerate(parsed.paragraphs):
        if not value.strip():
            continue
        segments.append(SourceSegment(f"SRC-{len(segments) + 1:03d}", parsed.document_version_id, parsed.locator(index), value))
    if not segments:
        raise ValueError("SOURCE_TEXT_UNAVAILABLE: no non-empty segments")
    return tuple(segments)


def resolve_segment(segments: tuple[SourceSegment, ...], segment_id: str, document_version_id: str) -> SourceSegment:
    for segment in segments:
        if segment.segment_id == segment_id:
            if segment.document_version_id != document_version_id:
                raise ValueError("segment belongs to another document version")
            return segment
    raise ValueError("unknown evidence segment")
